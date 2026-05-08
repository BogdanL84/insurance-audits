"""
pdf_extractor.py — Extract text from documents for AI analysis.

Supported types and how they're handled:
  PDF           → PyMuPDF, preserves page numbers
                  ↳ OCR fallback (Tesseract) when PyMuPDF returns <50 words
                    across all pages (i.e. a scanned image-only PDF)
  DOCX / DOC    → python-docx, includes table cells, ~3k chars per page
  TXT / MD      → read directly, split into ~3k-char pages
  XLSX / XLS    → openpyxl, extracts all cell values sheet by sheet
  PNG/JPG/JPEG  → NOT extracted — stored as reference files only

Returns: {page_num (int): text (str)}  — all types, 1-indexed
Use extract_with_info() to also receive {"method": "pdf_text"|"ocr"|"ocr_failed"|"empty", ...}.
"""

import io
import re
from pathlib import Path

# ── OCR config ─────────────────────────────────────────────────────
_TESSERACT_PATH      = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
_OCR_THRESHOLD_WORDS = 50   # PyMuPDF total < this triggers OCR fallback
_OCR_DPI             = 300  # rasterization DPI for scanned-page OCR

# ── File type categories ───────────────────────────────────────────
EXTRACTABLE_TYPES = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".md"}
IMAGE_TYPES       = {".png", ".jpg", ".jpeg"}
ALL_ACCEPTED      = EXTRACTABLE_TYPES | IMAGE_TYPES


def is_image(path: Path) -> bool:
    return path.suffix.lower() in IMAGE_TYPES


def is_extractable(path: Path) -> bool:
    return path.suffix.lower() in EXTRACTABLE_TYPES


# ── Main extraction entry point ────────────────────────────────────
def extract(path: Path) -> dict:
    """
    Extract text from a document. Returns {page_num: text}.
    Raises ValueError for images (call is_image() first).

    For OCR / extraction-method awareness, use extract_with_info().
    """
    return extract_with_info(path)[0]


def extract_with_info(path: Path) -> tuple[dict, dict]:
    """
    Like extract(), but also returns an info dict:
      {"method": "pdf_text" | "ocr" | "ocr_failed" | "empty"
                | "docx" | "xlsx" | "text",
       "pre_ocr_words": int (PDFs only, when method == 'ocr'),
       "ocr_pages_succeeded": int (PDFs, when method in {ocr, ocr_failed}),
       "ocr_pages_failed":    int (ditto)}
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _extract_docx(path), {"method": "docx"}
    elif suffix in (".xlsx", ".xls"):
        return _extract_xlsx(path), {"method": "xlsx"}
    elif suffix in (".txt", ".md"):
        return _extract_text(path), {"method": "text"}
    elif suffix in IMAGE_TYPES:
        raise ValueError(
            f"{path.name} is an image file. Images are stored as reference files, "
            "not extracted for text analysis."
        )
    else:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported: {', '.join(sorted(EXTRACTABLE_TYPES))}"
        )


# ── PDF ────────────────────────────────────────────────────────────
def _extract_pdf(path: Path) -> tuple[dict, dict]:
    """
    PyMuPDF extraction with OCR fallback for scanned image-only PDFs.

    Returns (text_by_page, info). info["method"] is:
      "pdf_text"   — PyMuPDF returned ≥ _OCR_THRESHOLD_WORDS total
      "ocr"        — fell back to Tesseract, OCR succeeded on ≥1 page
      "ocr_failed" — fell back to Tesseract, OCR returned no usable text
      "empty"      — PDF has zero pages or zero extractable content of any kind
    """
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

    result = {}
    doc = fitz.open(str(path))
    page_count = len(doc)
    try:
        for i in range(page_count):
            text = doc[i].get_text()
            text = re.sub(r"\n{3,}", "\n\n", text)
            text = re.sub(r" {2,}", " ", text)
            result[i + 1] = text.strip()
    finally:
        doc.close()

    if page_count == 0:
        return {1: ""}, {"method": "empty"}

    pre_ocr_words = sum(len(v.split()) for v in result.values())
    if pre_ocr_words >= _OCR_THRESHOLD_WORDS:
        return result, {"method": "pdf_text"}

    # PyMuPDF couldn't get meaningful text — try OCR.
    ocr_result, ocr_info = _ocr_pdf(path)
    ocr_words = sum(len(v.split()) for v in ocr_result.values())

    if ocr_words >= _OCR_THRESHOLD_WORDS:
        info = {
            "method":              "ocr",
            "pre_ocr_words":       pre_ocr_words,
            "ocr_pages_succeeded": ocr_info["pages_succeeded"],
            "ocr_pages_failed":    ocr_info["pages_failed"],
        }
        return ocr_result, info

    # OCR also produced nothing usable — return PyMuPDF's result so caller has
    # the page-count shape but flag it as ocr_failed for the UI.
    info = {
        "method":              "ocr_failed",
        "pre_ocr_words":       pre_ocr_words,
        "ocr_pages_succeeded": ocr_info["pages_succeeded"],
        "ocr_pages_failed":    ocr_info["pages_failed"],
    }
    return result, info


# ── OCR helper ─────────────────────────────────────────────────────
def _ocr_pdf(path: Path) -> tuple[dict, dict]:
    """
    Rasterize each page of a PDF at _OCR_DPI and run Tesseract on each image.
    Returns ({page_num: text}, {"pages_succeeded": N, "pages_failed": N}).
    Per-page failures are caught — one bad page never aborts the whole OCR pass.
    """
    import fitz
    import warnings
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        import pytesseract  # noqa: pulls pandas → noisy numpy 2.x ABI warnings
    from PIL import Image

    pytesseract.pytesseract.tesseract_cmd = _TESSERACT_PATH

    result   = {}
    succeeded = 0
    failed    = 0
    doc = fitz.open(str(path))
    try:
        for i in range(len(doc)):
            page = doc[i]
            try:
                pix       = page.get_pixmap(dpi=_OCR_DPI)
                img       = Image.open(io.BytesIO(pix.tobytes("png")))
                text      = pytesseract.image_to_string(img)
                text      = re.sub(r"\n{3,}", "\n\n", text)
                text      = re.sub(r" {2,}", " ", text).strip()
                result[i + 1] = text
                if len(text.split()) >= 3:
                    succeeded += 1
                else:
                    failed += 1
            except Exception as e:
                result[i + 1] = ""
                failed += 1
    finally:
        doc.close()
    return result, {"pages_succeeded": succeeded, "pages_failed": failed}


# ── DOCX ───────────────────────────────────────────────────────────
def _extract_docx(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc        = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Pull text from tables — contracts often put requirements in tables
    seen = set(paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in seen:
                    paragraphs.append(t)
                    seen.add(t)

    if not paragraphs:
        return {1: ""}

    return _chunk_text("\n\n".join(paragraphs))


# ── XLSX / XLS ─────────────────────────────────────────────────────
def _extract_xlsx(path: Path) -> dict:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")

    wb   = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    rows = []
    for sheet in wb.worksheets:
        rows.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))
    wb.close()

    if not rows:
        return {1: ""}
    return _chunk_text("\n".join(rows))


# ── TXT / MD ───────────────────────────────────────────────────────
def _extract_text(path: Path) -> dict:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError as e:
        raise OSError(f"Could not read {path.name}: {e}")

    if not text.strip():
        return {1: ""}
    return _chunk_text(text)


# ── Chunk helper ───────────────────────────────────────────────────
def _chunk_text(text: str, chunk_size: int = 3000) -> dict:
    """Split a flat string into ~3,000-char page chunks."""
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    return {i + 1: c for i, c in enumerate(chunks)} if chunks else {1: ""}


# ── Metadata helpers ───────────────────────────────────────────────
def get_page_count(path: Path) -> int:
    """Return page/chunk count without storing full text."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            import fitz
            doc   = fitz.open(str(path))
            count = len(doc)
            doc.close()
            return count
        elif suffix in IMAGE_TYPES:
            return 0  # images have no pages
        else:
            # Extract and count chunks for text-based formats
            extracted = extract(path)
            return len(extracted)
    except Exception:
        return 0


def format_size(path: Path) -> str:
    """Human-readable file size: '24 KB', '1.4 MB'."""
    try:
        size = path.stat().st_size
    except OSError:
        return "—"
    if size < 1_024:
        return f"{size} B"
    elif size < 1_024 * 1_024:
        return f"{size / 1_024:.0f} KB"
    else:
        return f"{size / (1_024 * 1_024):.1f} MB"


def get_file_size_mb(path: Path) -> float:
    try:
        return path.stat().st_size / (1024 * 1024)
    except OSError:
        return 0.0


def get_word_count(extracted: dict) -> int:
    return len(" ".join(extracted.values()).split())


# ── Save extracted text ────────────────────────────────────────────
def save_extracted_text(extracted: dict, output_path: Path) -> None:
    """Save with PAGE X of Y markers for easy reference during analysis."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    total = max(extracted.keys()) if extracted else 0
    lines = []
    for page_num in sorted(extracted.keys()):
        lines.append(f"\n{'='*60}")
        lines.append(f"PAGE {page_num} of {total}")
        lines.append(f"{'='*60}\n")
        lines.append(extracted[page_num])
    output_path.write_text("\n".join(lines), encoding="utf-8")


def build_extraction_preview(extracted: dict, max_chars: int = 400) -> str:
    if not extracted:
        return "(no text extracted)"
    text    = extracted.get(1, "")
    preview = text[:max_chars].strip()
    return preview + "…" if len(text) > max_chars else preview
